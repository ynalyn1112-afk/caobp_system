from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
from django.http import JsonResponse, HttpResponse
from django.contrib import messages
from django.db.models import Sum, Count, Q
from django.utils import timezone
from django.core.paginator import Paginator
from django.contrib.auth import get_user_model
from datetime import datetime, timedelta
from .models import Notification, OPBRequest, OPBItem
from accounts.models import User
import json
import csv
import os
import shutil
import zipfile
from io import StringIO, BytesIO
from django.conf import settings
from django.core.management import call_command
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from decimal import Decimal

def decimal_default(obj):
    if isinstance(obj, Decimal):
        return float(obj)
    raise TypeError

@login_required
@staff_member_required
def admin_dashboard(request):
    # Get statistics
    total_departments = User.objects.filter(role='unit_head').values('department').distinct().count()
    total_opb_requests = OPBRequest.objects.count()
    total_requests = total_opb_requests 
    pending_requests = OPBRequest.objects.filter(status='pending').count()
    
    # Get all departments
    all_departments = User.objects.filter(role='unit_head').values_list('department', flat=True).distinct()
    
    # Calculate submitted departments
    submitted_depts = set()
    submitted_breakdown = []
    for dept in all_departments:
        dept_count = OPBRequest.objects.filter(department=dept).count()
        if dept_count > 0:
            submitted_depts.add(dept)
            submitted_breakdown.append({
                'department': dept,
                'count': dept_count,
                'display': dict(User.UNIT_CHOICES).get(dept, dept)
            })
    
    # Calculate not submitted departments
    not_submitted_depts = set(all_departments) - submitted_depts
    not_submitted_breakdown = []
    for dept in not_submitted_depts:
        not_submitted_breakdown.append({
            'department': dept,
            'display': dict(User.UNIT_CHOICES).get(dept, dept)
        })
    
    total_submitted = len(submitted_depts)
    total_not_submitted = len(not_submitted_depts)
    percent_submitted = (total_submitted / total_departments * 100) if total_departments > 0 else 0
    
    # Get recent requests
    recent_opb = OPBRequest.objects.filter(status='pending').order_by('-created_at')[:5]
    
    # Get budget allocation by department (combine PRE, and OPB)
    departments = User.objects.filter(role='unit_head').values_list('department', flat=True).distinct()
    dept_budget_data = []
    
    for dept in departments:
        
        opb_total = sum([req.total_budget_amount for req in OPBRequest.objects.filter(department=dept, status='for-approval')])
        combined_total = opb_total
        
        if combined_total > 0:  # Only include departments with budget
            dept_budget_data.append({
                'department': dept,
                'total': combined_total,
                'opb': opb_total
            })
    
    # If no real data, provide sample data for demonstration
    if not dept_budget_data:
         dept_budget_data = [
             {'department': 'EDUC', 'total': 500000, 'opb': 500000},
             {'department': 'ENGINEERING', 'total': 400000, 'opb': 400000},
             {'department': 'IT', 'total': 350000, 'opb': 350000},
             {'department': 'CAS', 'total': 300000, 'opb': 300000},
             {'department': 'AGRI', 'total': 250000, 'opb': 250000},
         ]
    
    # Sort by total budget descending
    dept_budget_data.sort(key=lambda x: x['total'], reverse=True)
    
    # Get monthly budget data for the line chart (last 12 months)
    from datetime import datetime, timedelta
    monthly_data = []
    for i in range(12):
        month_start = datetime.now().replace(day=1) - timedelta(days=30*i)
        month_end = month_start + timedelta(days=30)
        
        month_opb = sum([req.total_budget_amount for req in OPBRequest.objects.filter(
            created_at__date__gte=month_start.date(),
            created_at__date__lt=month_end.date()
        )])
        
        monthly_data.append({
            'month': month_start.strftime('%b'),
            'total': month_opb
        })
    
    monthly_data.reverse()  # Show oldest to newest
    
    # If no real monthly data, provide sample data
    if all(month['total'] == 0 for month in monthly_data):
         import random
         monthly_data = []
         months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
         for i, month in enumerate(months):
             monthly_data.append({
                 'month': month,
                 'total': random.randint(100000, 500000)  # Sample budget data
             })

    context = {
        'total_departments': total_departments,
        'total_opb_requests': total_opb_requests,
        'total_requests': total_requests,
        'pending_requests': pending_requests,
        'total_submitted': total_submitted,
        'total_not_submitted': total_not_submitted,
        'percent_submitted': round(percent_submitted, 1),
        'submitted_breakdown': submitted_breakdown,
        'not_submitted_breakdown': not_submitted_breakdown,
        'recent_opb': recent_opb,
        'dept_budget_data': json.dumps(dept_budget_data, default=decimal_default),
        'monthly_data': json.dumps(monthly_data, default=decimal_default),
    }
    
    return render(request, 'admin_dashboard.html', context)


@login_required
@staff_member_required
def admin_users(request):
    # Get search query
    search_query = request.GET.get('search', '')
    users = User.objects.all()
    
    if search_query:
        users = users.filter(
            Q(username__icontains=search_query) |
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(email__icontains=search_query)
        )
    
    users = users.order_by('-date_joined')
    
    # Pagination
    paginator = Paginator(users, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Get departments that already have users
    taken_departments = set(User.objects.filter(department__isnull=False).values_list('department', flat=True))
    
    context = {
        'users': page_obj,
        'search_query': search_query,
        'unit_choices': User.UNIT_CHOICES,
        'taken_departments': json.dumps(list(taken_departments)),
    }
    
    return render(request, 'admin_users.html', context)

@login_required
@staff_member_required
def admin_reports(request):
    # Get filter parameters
    dept_filter = request.GET.get('department', '')
    status_filter = request.GET.get('status', '')
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')
    export_format = request.GET.get('export', '')
    
    # Get data for reports
    opb_requests = OPBRequest.objects.select_related('department_head').all()
    
    if status_filter:
        opb_requests = opb_requests.filter(status=status_filter)

    if dept_filter:
        opb_requests = opb_requests.filter(department=dept_filter)
    
    if date_from:
        opb_requests = opb_requests.filter(created_at__date__gte=date_from)
    
    if date_to:
        opb_requests = opb_requests.filter(created_at__date__lte=date_to)
    
    # Calculate totals
    opb_total = sum(request.total_budget_amount for request in opb_requests)
    
    # Calculate summary statistics
    total_requests = opb_requests.count()
    approved_requests = opb_requests.filter(status='for-approval').count()
    pending_requests = opb_requests.filter(status='pending').count()
    rejected_requests = opb_requests.filter(status='enhancement').count()
    
    # Department ranking
    dept_ranking = []
    departments = User.objects.filter(role='unit_head').values_list('department', flat=True).distinct()
    for dept in departments:
        
        dept_opb = sum(request.total_budget_amount for request in opb_requests.filter(department=dept))
        dept_total = dept_opb 
        dept_ranking.append({
            'department': dept,
            'total': dept_total,
            'opb': dept_opb,
           
        })
    dept_ranking.sort(key=lambda x: x['total'], reverse=True)
    
    # Handle export
    if export_format == 'csv':
        return export_reports_csv(opb_requests, dept_ranking)
    elif export_format == 'pdf':
        return export_reports_pdf(opb_requests, dept_ranking)
    elif export_format == 'docx':
        return export_reports_docx(opb_requests, dept_ranking)
    elif export_format == 'print':
        return render(request, 'admin_reports_print.html', {
            'opb_requests': opb_requests,
            'dept_ranking': dept_ranking,
            'total_requests': total_requests,
            'approved_requests': approved_requests,
            'pending_requests': pending_requests,
            'rejected_requests': rejected_requests,
            'opb_total': opb_total,
        })
    
    context = {
        'opb_requests': opb_requests[:10],  # Limit for display
        'unit_choices': User.UNIT_CHOICES,
        'opb_total': opb_total,
        'total_requests': total_requests,
        'approved_requests': approved_requests,
        'pending_requests': pending_requests,
        'rejected_requests': rejected_requests,
        'dept_ranking': dept_ranking,
        'dept_filter': dept_filter,
        'date_from': date_from,
        'date_to': date_to,
        'status_filter': status_filter,
    }
    
    return render(request, 'admin_reports.html', context)


@login_required
@staff_member_required
def admin_settings(request):
    # Get system statistics
    total_users = User.objects.count()
    total_notifications = Notification.objects.count()
    all_departments = User.objects.filter(role='unit_head').values_list('department', flat=True).distinct()
    
    
    # Get database size (approximate)
    import os
    db_path = settings.DATABASES['default']['NAME']
    db_size = 0
    if os.path.exists(db_path):
        db_size = os.path.getsize(db_path)
    
    # Get system health metrics
    active_users = User.objects.filter(is_active=True).count()
    inactive_users = User.objects.filter(is_active=False).count()
    submitted_depts = set()
    submitted_breakdown = []
    for dept in all_departments:
        dept_count = OPBRequest.objects.filter(department=dept).count()
        if dept_count > 0:
            submitted_depts.add(dept)
            submitted_breakdown.append({
                'department': dept,
                'count': dept_count,
                'display': dict(User.UNIT_CHOICES).get(dept, dept)
            })
    total_submitted = len(submitted_depts)
    # Get backup information
    backup_dir = os.path.join(settings.BASE_DIR, 'backups')
    backup_files = []
    if os.path.exists(backup_dir):
        backup_files = [f for f in os.listdir(backup_dir) if f.endswith('.sqlite3')]
        backup_files.sort(key=lambda x: os.path.getmtime(os.path.join(backup_dir, x)), reverse=True)
    
    context = {
        'total_users': total_users,
        'total_notifications': total_notifications,
        'db_size': db_size,
        'active_users': active_users,
        'inactive_users': inactive_users,
        'total_opb_submitted': total_submitted,
        'backup_files': backup_files,
    }
    
    return render(request, 'admin_settings.html', context)


# AJAX Views for CRUD operations


@login_required
@staff_member_required
def ajax_add_user(request):
    """AJAX endpoint to add user"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            
            # Check if username already exists
            if User.objects.filter(username=data['username']).exists():
                return JsonResponse({
                    'success': False,
                    'message': 'Username already exists'
                })
            
            # Check if email already exists
            if User.objects.filter(email=data['email']).exists():
                return JsonResponse({
                    'success': False,
                    'message': 'Email already exists'
                })
            
            # Check if department already has a user (only for unit_head role)
            if data['role'] == 'unit_head' and data.get('department'):
                existing_user = User.objects.filter(department=data['department']).first()
                if existing_user:
                    department_display = dict(User.UNIT_CHOICES).get(data['department'], data['department'])
                    return JsonResponse({
                        'success': False,
                        'message': f"There's already an existing user in this Unit ({department_display})"
                    })
            
            # Create user
            user = User.objects.create_user(
                username=data['username'],
                email=data['email'],
                password=data['password'],
                first_name=data['first_name'],
                last_name=data['last_name'],
                department=data.get('department', ''),
                role=data['role']
            )
            
            return JsonResponse({
                'success': True,
                'message': 'User created successfully',
                'user_id': user.id
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
@staff_member_required
def ajax_edit_user(request):
    """AJAX endpoint to edit user"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            user_id = data['user_id']
            
            user = get_object_or_404(User, id=user_id)
            
            # Check if department is already taken by another user
            if data['role'] == 'unit_head' and data.get('department'):
                existing_user = User.objects.filter(department=data['department']).exclude(id=user_id).first()
                if existing_user:
                    department_display = dict(User.UNIT_CHOICES).get(data['department'], data['department'])
                    return JsonResponse({
                        'success': False,
                        'message': f"There's already an existing user in this Unit ({department_display})"
                    })
            
            # Update user fields
            user.first_name = data['first_name']
            user.last_name = data['last_name']
            user.email = data['email']
            user.department = data.get('department', '')
            user.role = data['role']
            
            # Update password if provided
            if data.get('password'):
                user.set_password(data['password'])
            
            user.save()
            
            return JsonResponse({
                'success': True,
                'message': 'User updated successfully'
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
@staff_member_required
def ajax_delete_user(request):
    """AJAX endpoint to delete user"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            user_id = data['user_id']
            
            user = get_object_or_404(User, id=user_id)
            
            # Don't allow deleting the current user
            if user == request.user:
                return JsonResponse({
                    'success': False,
                    'message': 'Cannot delete your own account'
                })
            
            user.delete()
            
            return JsonResponse({
                'success': True,
                'message': 'User deleted successfully'
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
@staff_member_required
def ajax_toggle_user_status(request):
    """AJAX endpoint to toggle user active status"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            user_id = data['user_id']
            
            user = get_object_or_404(User, id=user_id)
            
            # Don't allow deactivating the current user
            if user == request.user:
                return JsonResponse({
                    'success': False,
                    'message': 'Cannot deactivate your own account'
                })
            
            user.is_active = not user.is_active
            user.save()
            
            status = 'activated' if user.is_active else 'deactivated'
            
            return JsonResponse({
                'success': True,
                'message': f'User {status} successfully'
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


# Department Head
@login_required
def head_dashboard(request):
    # Get user's department
    user_dept = request.user.department
    
    # Get statistics for this department head
    
    submitted_count = 0 
    approved_count = 0
    rejected_count = 0
    pending_count = 0
    
    # Get recent notifications
    notifications = Notification.objects.filter(user=request.user).order_by('-created_at')[:5]
    unread_count = Notification.objects.filter(user=request.user, is_read=False).count()
    
    
    
    context = {
        'submitted_count': submitted_count,
        'approved_count': approved_count,
        'rejected_count': rejected_count,
        'pending_count': pending_count,
        'notifications': notifications,
        'unread_count': unread_count,
        'user_dept': user_dept,
    }
    
    return render(request, 'head_dashboard.html', context)

@login_required
def head_notifications(request):
    # Handle bulk actions
    if request.method == 'POST':
        action = request.POST.get('action')
        
        if action == 'mark_all_read':
            Notification.objects.filter(user=request.user, is_read=False).update(is_read=True)
            messages.success(request, 'All notifications marked as read!')
        elif action == 'delete_all':
            Notification.objects.filter(user=request.user).delete()
            messages.success(request, 'All notifications deleted!')
        
        return redirect('head_notifications')
    
    # Get all notifications for this user
    notifications = Notification.objects.filter(user=request.user).order_by('-created_at')
    
    context = {
        'notifications': notifications,
    }
    
    return render(request, 'head_notifications.html', context)

# OPB Views
@login_required
def head_opb_requests(request):
    # Handle form submission
    if request.method == 'POST':
        try:
            # Get common fields
            fiscal_year = request.POST.get('fiscal_year', '2026')
            unit = request.user.get_department_display() if request.user.department else 'No Unit Assigned'
            
            # Get array fields for multiple entries
            kra_nos = request.POST.getlist('kra_no[]') if 'kra_no[]' in request.POST else [request.POST.get('kra_no', '')]
            objective_nos = request.POST.getlist('objective_no[]') if 'objective_no[]' in request.POST else [request.POST.get('objective_no', '')]
            indicators_list = request.POST.getlist('indicators[]') if 'indicators[]' in request.POST else [request.POST.get('indicators', '')]
            annual_targets = request.POST.getlist('annual_target[]') if 'annual_target[]' in request.POST else [request.POST.get('annual_target', '')]
            activities_list = request.POST.getlist('activities[]') if 'activities[]' in request.POST else [request.POST.get('activities', '')]
            timeframes = request.POST.getlist('timeframe[]') if 'timeframe[]' in request.POST else [request.POST.get('timeframe', '')]
            budget_amounts = request.POST.getlist('budget_amount[]') if 'budget_amount[]' in request.POST else [request.POST.get('budget_amount', 0)]
            source_of_funds = request.POST.getlist('source_of_fund[]') if 'source_of_fund[]' in request.POST else [request.POST.get('source_of_fund', '')]
            responsible_units_list = request.POST.getlist('responsible_units[]') if 'responsible_units[]' in request.POST else [request.POST.get('responsible_units', '')]
            
            # Create single OPB request
            opb_request = OPBRequest.objects.create(
                department_head=request.user,
                department=request.user.department,
                fiscal_year=fiscal_year,
                unit=unit,
            )
            
            # Create OPB items for each row
            created_count = 0
            for i in range(len(kra_nos)):
                if kra_nos[i].strip() or objective_nos[i].strip() or indicators_list[i].strip():
                    OPBItem.objects.create(
                        request=opb_request,
                        kra_no=kra_nos[i] if i < len(kra_nos) else '',
                        objective_no=objective_nos[i] if i < len(objective_nos) else '',
                        indicators=indicators_list[i] if i < len(indicators_list) else '',
                        annual_target=annual_targets[i] if i < len(annual_targets) else '',
                        activities=activities_list[i] if i < len(activities_list) else '',
                        timeframe=timeframes[i] if i < len(timeframes) else '',
                        budget_amount=float(budget_amounts[i]) if i < len(budget_amounts) and budget_amounts[i] else 0,
                        source_of_fund=source_of_funds[i] if i < len(source_of_funds) else '',
                        responsible_units=responsible_units_list[i] if i < len(responsible_units_list) else '',
                    )
                    created_count += 1
            
            if created_count == 0:
                opb_request.delete()  # Delete the empty request
                messages.error(request, 'Please fill in at least one row with data')
            else:
                messages.success(request, f'OPB request submitted successfully! ({created_count} entries)')
            return redirect('head_opb')
        except Exception as e:
            messages.error(request, f'Error submitting request: {str(e)}')
    
    # Get user's OPB requests
    requests = OPBRequest.objects.filter(department_head=request.user).order_by('-created_at')
    
    # Get user's department display name
    user_department_display = request.user.get_department_display() if request.user.department else 'No Unit Assigned'
    unread_count = Notification.objects.filter(user=request.user, is_read=False).count()
    
    context = {
        'requests': requests,
        'user_department': request.user.department,
        'user_department_display': user_department_display,
        'unread_count': unread_count,
    }
    
    return render(request, 'head_opb.html', context)


@login_required
def head_opb_edit(request, request_id):
    opb_request = get_object_or_404(OPBRequest, id=request_id, department_head=request.user)
    
    if request.method == 'POST':
        try:
            # Update basic request info
            opb_request.fiscal_year = request.POST.get('fiscal_year', '2026')
            opb_request.unit = request.user.get_department_display() if request.user.department else 'No Unit Assigned'
            opb_request.status = 'pending'  # Reset status to pending after edit
            opb_request.save()
            
            # Clear existing items
            opb_request.items.all().delete()
            
            # Get array fields for multiple entries
            kra_nos = request.POST.getlist('kra_no[]')
            objective_nos = request.POST.getlist('objective_no[]')
            indicators_list = request.POST.getlist('indicators[]')
            annual_targets = request.POST.getlist('annual_target[]')
            activities_list = request.POST.getlist('activities[]')
            timeframes = request.POST.getlist('timeframe[]')
            budget_amounts = request.POST.getlist('budget_amount[]')
            source_of_funds = request.POST.getlist('source_of_fund[]')
            responsible_units_list = request.POST.getlist('responsible_units[]')
            
            # Create new OPB items
            created_count = 0
            for i in range(len(kra_nos)):
                if kra_nos[i].strip() or objective_nos[i].strip() or indicators_list[i].strip():
                    OPBItem.objects.create(
                        request=opb_request,
                        kra_no=kra_nos[i] if i < len(kra_nos) else '',
                        objective_no=objective_nos[i] if i < len(objective_nos) else '',
                        indicators=indicators_list[i] if i < len(indicators_list) else '',
                        annual_target=annual_targets[i] if i < len(annual_targets) else '',
                        activities=activities_list[i] if i < len(activities_list) else '',
                        timeframe=timeframes[i] if i < len(timeframes) else '',
                        budget_amount=float(budget_amounts[i]) if i < len(budget_amounts) and budget_amounts[i] else 0,
                        source_of_fund=source_of_funds[i] if i < len(source_of_funds) else '',
                        responsible_units=responsible_units_list[i] if i < len(responsible_units_list) else '',
                    )
                    created_count += 1
            
            if created_count == 0:
                messages.error(request, 'Please fill in at least one row with data')
            else:
                messages.success(request, f'OPB request updated successfully! ({created_count} entries)')
            return redirect('head_opb')
        except Exception as e:
            messages.error(request, f'Error updating request: {str(e)}')
    
    # Get user's department display name
    user_department_display = request.user.get_department_display() if request.user.department else 'No Unit Assigned'
    unread_count = Notification.objects.filter(user=request.user, is_read=False).count()
    
    context = {
        'opb_request': opb_request,
        'user_department': request.user.department,
        'user_department_display': user_department_display,
        'unread_count': unread_count,
    }
    return render(request, 'head_opb_edit.html', context)


@login_required
def head_opb_view(request, request_id):
    opb_request = get_object_or_404(OPBRequest, id=request_id, department_head=request.user)
    unread_count = Notification.objects.filter(user=request.user, is_read=False).count()
    
    context = {
        'opb_request': opb_request,
        'unread_count': unread_count,
    }
    return render(request, 'head_opb_view.html', context)


@login_required
@staff_member_required
def admin_opb_requests(request):
    # Get search and filter parameters
    search_query = request.GET.get('search', '')
    status_filter = request.GET.get('status', '')
    dept_filter = request.GET.get('department', '')
    
    requests = OPBRequest.objects.select_related('department_head').all()
    
    if search_query:
        requests = requests.filter(
            Q(items__kra_no__icontains=search_query) |
            Q(items__indicators__icontains=search_query) |
            Q(department_head__first_name__icontains=search_query) |
            Q(department_head__last_name__icontains=search_query) |
            Q(unit__icontains=search_query)
        ).distinct()
    
    if status_filter:
        requests = requests.filter(status=status_filter)
    
    if dept_filter:
        requests = requests.filter(department=dept_filter)
    
    requests = requests.order_by('-created_at')
    
    # Pagination
    paginator = Paginator(requests, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'requests': page_obj,
        'search_query': search_query,
        'status_filter': status_filter,
        'dept_filter': dept_filter,
        'unit_choices': User.UNIT_CHOICES,
    }
    
    return render(request, 'admin_opb.html', context)


@login_required
@staff_member_required
def admin_opb_view_details(request, request_id):
    opb_request = get_object_or_404(OPBRequest, id=request_id)
    
    context = {
        'request': opb_request,
    }
    
    return render(request, 'admin_opb_details.html', context)


# AJAX Views for Department Head

@login_required
def ajax_mark_notification_read(request):
    """AJAX endpoint to mark notification as read"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            notification_id = data['notification_id']
            
            notification = get_object_or_404(Notification, id=notification_id, user=request.user)
            notification.is_read = True
            notification.save()
            
            return JsonResponse({
                'success': True,
                'message': 'Notification marked as read'
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
def ajax_delete_notification(request):
    """AJAX endpoint to delete notification"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            notification_id = data['notification_id']
            
            notification = get_object_or_404(Notification, id=notification_id, user=request.user)
            notification.delete()
            
            return JsonResponse({
                'success': True,
                'message': 'Notification deleted'
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})

@login_required
def ajax_mark_all_notifications_read(request):
    """AJAX endpoint to mark all notifications as read"""
    if request.method == 'POST':
        try:
            Notification.objects.filter(user=request.user, is_read=False).update(is_read=True)
            
            return JsonResponse({
                'success': True,
                'message': 'All notifications marked as read'
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
def ajax_delete_all_notifications(request):
    """AJAX endpoint to delete all notifications"""
    if request.method == 'POST':
        try:
            Notification.objects.filter(user=request.user).delete()
            
            return JsonResponse({
                'success': True,
                'message': 'All notifications deleted'
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


# OPB AJAX Views
@login_required
def ajax_submit_opb_request(request):
    """AJAX endpoint to submit OPB request"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            
            # Get common fields
            fiscal_year = data.get('fiscal_year', '2026')
            unit = data.get('unit', '')
            
            # Get array fields for multiple entries
            kra_nos = data.getlist('kra_no[]') if 'kra_no[]' in data else [data.get('kra_no', '')]
            objective_nos = data.getlist('objective_no[]') if 'objective_no[]' in data else [data.get('objective_no', '')]
            indicators_list = data.getlist('indicators[]') if 'indicators[]' in data else [data.get('indicators', '')]
            annual_targets = data.getlist('annual_target[]') if 'annual_target[]' in data else [data.get('annual_target', '')]
            activities_list = data.getlist('activities[]') if 'activities[]' in data else [data.get('activities', '')]
            timeframes = data.getlist('timeframe[]') if 'timeframe[]' in data else [data.get('timeframe', '')]
            budget_amounts = data.getlist('budget_amount[]') if 'budget_amount[]' in data else [data.get('budget_amount', 0)]
            source_of_funds = data.getlist('source_of_fund[]') if 'source_of_fund[]' in data else [data.get('source_of_fund', '')]
            responsible_units_list = data.getlist('responsible_units[]') if 'responsible_units[]' in data else [data.get('responsible_units', '')]
            
            # Create single OPB request
            opb_request = OPBRequest.objects.create(
                department_head=request.user,
                department=request.user.department,
                fiscal_year=fiscal_year,
                unit=unit,
            )
            
            # Create OPB items for each row
            created_count = 0
            for i in range(len(kra_nos)):
                if kra_nos[i].strip() or objective_nos[i].strip() or indicators_list[i].strip():
                    OPBItem.objects.create(
                        request=opb_request,
                        kra_no=kra_nos[i] if i < len(kra_nos) else '',
                        objective_no=objective_nos[i] if i < len(objective_nos) else '',
                        indicators=indicators_list[i] if i < len(indicators_list) else '',
                        annual_target=annual_targets[i] if i < len(annual_targets) else '',
                        activities=activities_list[i] if i < len(activities_list) else '',
                        timeframe=timeframes[i] if i < len(timeframes) else '',
                        budget_amount=float(budget_amounts[i]) if i < len(budget_amounts) and budget_amounts[i] else 0,
                        source_of_fund=source_of_funds[i] if i < len(source_of_funds) else '',
                        responsible_units=responsible_units_list[i] if i < len(responsible_units_list) else '',
                    )
                    created_count += 1
            
            if created_count == 0:
                opb_request.delete()  # Delete the empty request
                return JsonResponse({
                    'success': False,
                    'message': 'Please fill in at least one row with data'
                })
            
            return JsonResponse({
                'success': True,
                'message': f'OPB request submitted successfully ({created_count} entries)',
                'request_id': str(opb_request.id)
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
def ajax_delete_opb_request(request):
    """AJAX endpoint to delete OPB request"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            request_id = data['request_id']
            
            opb_request = get_object_or_404(OPBRequest, id=request_id, department_head=request.user)
            opb_request.delete()
            
            return JsonResponse({
                'success': True,
                'message': 'OPB request deleted successfully'
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
@staff_member_required
def ajax_approve_request(request):
    """AJAX endpoint to approve budget request"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            request_type = data['type']
            request_id = data['request_id']
            
            
            if request_type == 'opb':
                req = get_object_or_404(OPBRequest, id=request_id)
            else:
                return JsonResponse({
                    'success': False,
                    'message': 'Invalid request type'
                })
            
            # Check if request is already processed
            if req.status != 'pending':
                return JsonResponse({
                    'success': False,
                    'message': 'Request has already been processed'
                })
            
            # Approve the request
            req.status = 'for-approval'
            req.admin_notes = data.get('notes', '')
            req.save()
            
            # Create notification
            if request_type == 'opb':
                request_description = f"OPB for {req.unit or 'your unit'}"
            else:
                request_description = f"{request_type.upper()} request for {req.program_title}"
            
            Notification.objects.create(
                user=req.department_head,
                title=f'OPB for {req.unit} For Approval',
                message=f'Your {request_description} has been set for approval.',
                notification_type='success'
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Request approved successfully'
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
@staff_member_required
def ajax_reject_request(request):
    """AJAX endpoint to reject budget request"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            request_type = data['type']
            request_id = data['request_id']
            
            
            if request_type == 'opb':
                req = get_object_or_404(OPBRequest, id=request_id)
            else:
                return JsonResponse({
                    'success': False,
                    'message': 'Invalid request type'
                })
            
            req.status = 'enhancement'
            req.admin_notes = data.get('notes', data.get('reason', ''))
            req.save()
            
            # Create notification
            if request_type == 'opb':
                request_description = f"OPB for {req.unit or 'your unit'}"
            else:
                request_description = f"{request_type.upper()} request for {req.program_title}"
            
            Notification.objects.create(
                user=req.department_head,
                title=f'OPB for {req.unit} For enhancement',
                message=f'Your {request_description} has been set for enhancement.',
                notification_type='warning'
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Request rejected successfully'
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
@staff_member_required
def ajax_get_user(request, user_id):
    """AJAX endpoint to get user data for editing"""
    if request.method == 'GET':
        try:
            user = get_object_or_404(User, id=user_id)
            
            return JsonResponse({
                'success': True,
                'user': {
                    'id': user.id,
                    'first_name': user.first_name,
                    'last_name': user.last_name,
                    'username': user.username,
                    'email': user.email,
                    'department': user.department,
                    'role': user.role,
                }
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
@staff_member_required
def ajax_create_backup(request):
    """AJAX endpoint to create database backup"""
    if request.method == 'POST':
        try:
            # Create backups directory if it doesn't exist
            backup_dir = os.path.join(settings.BASE_DIR, 'backups')
            os.makedirs(backup_dir, exist_ok=True)
            
            # Generate backup filename with timestamp
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            backup_filename = f'caobp_backup_{timestamp}.sqlite3'
            backup_path = os.path.join(backup_dir, backup_filename)
            
            # Get current database path
            db_path = settings.DATABASES['default']['NAME']
            
            if not os.path.exists(db_path):
                return JsonResponse({
                    'success': False,
                    'message': 'Database file not found'
                })
            
            # Copy database file to backup location
            shutil.copy2(db_path, backup_path)
            
            # Create metadata file
            metadata = {
                'created_at': datetime.now().isoformat(),
                'created_by': request.user.username,
                'database_size': os.path.getsize(db_path),
                'total_users': User.objects.count(),
                'total_opb_requests': OPBRequest.objects.count(),
            }
            
            # Save metadata as JSON file
            metadata_filename = f'caobp_backup_{timestamp}_metadata.json'
            metadata_path = os.path.join(backup_dir, metadata_filename)
            with open(metadata_path, 'w') as f:
                json.dump(metadata, f, indent=2)
            
            return JsonResponse({
                'success': True,
                'message': f'Database backup created successfully: {backup_filename}',
                'backup_filename': backup_filename
            })
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': f'Failed to create backup: {str(e)}'
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
@staff_member_required
def ajax_download_backup(request):
    """AJAX endpoint to download backup file"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            backup_filename = data['filename']
            
            backup_dir = os.path.join(settings.BASE_DIR, 'backups')
            backup_path = os.path.join(backup_dir, backup_filename)
            
            if not os.path.exists(backup_path):
                return JsonResponse({
                    'success': False,
                    'message': 'Backup file not found'
                })
            
            with open(backup_path, 'rb') as f:
                response = HttpResponse(f.read(), content_type='application/octet-stream')
                response['Content-Disposition'] = f'attachment; filename="{backup_filename}"'
                return response
                
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': f'Failed to download backup: {str(e)}'
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
@staff_member_required
def ajax_delete_backup(request):
    """AJAX endpoint to delete backup file"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            backup_filename = data['filename']
            base, _ = os.path.splitext(backup_filename)
            metadata_filename = f'{base}_metadata.json'
            
            backup_dir = os.path.join(settings.BASE_DIR, 'backups')
            backup_path = os.path.join(backup_dir, backup_filename)
            metadata_path = os.path.join(backup_dir, metadata_filename)
            
            deleted_any = False

            # Delete main backup
            if os.path.exists(backup_path):
                os.remove(backup_path)
                deleted_any = True

            # Delete metadata
            if os.path.exists(metadata_path):
                os.remove(metadata_path)
                deleted_any = True

            if deleted_any:
                return JsonResponse({
                    'success': True,
                    'message': 'Backup deleted successfully'
                })
            else:
                return JsonResponse({
                    'success': False,
                    'message': 'Backup file not found'
                })
                
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': f'Failed to delete backup: {str(e)}'
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


@login_required
@staff_member_required
def ajax_restore_backup(request):
    """AJAX endpoint to restore from backup"""
    if request.method == 'POST':
        try:
            # Handle both FormData and JSON
            if request.content_type == 'application/json':
                data = json.loads(request.body)
            else:
                data = request.POST
            backup_filename = data['filename']
            
            backup_dir = os.path.join(settings.BASE_DIR, 'backups')
            backup_path = os.path.join(backup_dir, backup_filename)
            
            if not os.path.exists(backup_path):
                return JsonResponse({
                    'success': False,
                    'message': 'Backup file not found'
                })
            
            # Get current database path
            db_path = settings.DATABASES['default']['NAME']
            
            # Create a backup of current database before restoring
            current_backup = f"{db_path}.backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            if os.path.exists(db_path):
                shutil.copy2(db_path, current_backup)
            
            # Replace current database with backup
            shutil.copy2(backup_path, db_path)
            
            return JsonResponse({
                'success': True,
                'message': 'Database restored successfully. The system is now using the backup database.'
            })
                
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': f'Failed to restore backup: {str(e)}'
            })
    return JsonResponse({'success': False, 'message': 'Invalid request'})


# Export Functions
def export_reports_csv(opb_requests, dept_ranking):
    """Export reports as CSV"""
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="caobp_report.csv"'
    
    writer = csv.writer(response)
    
    # Write header with timestamp
    writer.writerow(['CAOBP SYSTEM - REPORT'])
    writer.writerow(['Generated on:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')])
    writer.writerow(['Generated by:', 'System Administrator'])
    writer.writerow([])
    
    # Calculate comprehensive statistics
    total_requests = opb_requests.count()
    approved_requests = opb_requests.filter(status='for-approval').count()
    pending_requests = opb_requests.filter(status='pending').count()
    rejected_requests = opb_requests.filter(status='enhancement').count()
    
    total_budget_amount = sum(r.total_budget_amount for r in opb_requests)
    approved_budget_amount = sum(r.total_budget_amount for r in opb_requests.filter(status='for-approval'))
    pending_budget_amount = sum(r.total_budget_amount for r in opb_requests.filter(status='pending'))
    rejected_budget_amount = sum(r.total_budget_amount for r in opb_requests.filter(status='enhancement'))
    
    # Get all departments and calculate submission statistics
    all_departments = User.objects.filter(role='unit_head').values_list('department', flat=True).distinct()
    submitted_departments = set(opb_requests.values_list('department', flat=True))
    not_submitted_departments = set(all_departments) - submitted_departments
    
    submission_percentage = (len(submitted_departments) / len(all_departments) * 100) if all_departments else 0
    
    # EXECUTIVE SUMMARY SECTION
    writer.writerow(['=' * 80])
    writer.writerow(['OVERALL SUMMARY'])
    writer.writerow(['=' * 80])
    writer.writerow([])
    
    writer.writerow(['Overall Statistics'])
    writer.writerow(['Total Units', len(all_departments)])
    writer.writerow(['Units with Submissions', len(submitted_departments)])
    writer.writerow(['Units without Submissions', len(not_submitted_departments)])
    writer.writerow(['Submission Rate', f"{submission_percentage:.1f}%"])
    writer.writerow([])
    
    writer.writerow(['OPB SUBMISSION STATISTICS'])
    writer.writerow(['Total OPB Requests', total_requests])
    writer.writerow(['For Approval', approved_requests])
    writer.writerow(['Pending', pending_requests])
    writer.writerow(['Enhancement', rejected_requests])
    writer.writerow([])
    
    writer.writerow(['BUDGET STATISTICS'])
    writer.writerow(['Total Budget Amount', f"₱{total_budget_amount:,.2f}"])
    writer.writerow(['For Approval Budget Amount', f"₱{approved_budget_amount:,.2f}"])
    writer.writerow(['Pending Budget Amount', f"₱{pending_budget_amount:,.2f}"])
    writer.writerow(['Enhancement Budget Amount', f"₱{rejected_budget_amount:,.2f}"])
    writer.writerow([])
    
    # DEPARTMENT SUBMISSION ANALYSIS
    writer.writerow(['=' * 80])
    writer.writerow(['Units SUBMISSION ANALYSIS'])
    writer.writerow(['=' * 80])
    writer.writerow([])
    
    writer.writerow(['Units with Submissions'])
    writer.writerow(['Unit', 'Display Name', 'Submissions', 'Total Budget', 'Status Distribution'])
    
    for dept in submitted_departments:
        dept_requests = opb_requests.filter(department=dept)
        dept_count = dept_requests.count()
        dept_total = sum(r.total_budget_amount for r in dept_requests)
        dept_approved = dept_requests.filter(status='for-approval').count()
        dept_pending = dept_requests.filter(status='pending').count()
        dept_rejected = dept_requests.filter(status='enhancement').count()
        
        dept_display = dict(User.UNIT_CHOICES).get(dept, dept)
        status_dist = f"A:{dept_approved} P:{dept_pending} R:{dept_rejected}"
        
        writer.writerow([
            dept,
            dept_display,
            dept_count,
            f"₱{dept_total:,.2f}",
            status_dist
        ])
    
    writer.writerow([])
    
    writer.writerow(['Units without Submissions'])
    writer.writerow(['Unit', 'Display Name', 'Status'])
    for dept in not_submitted_departments:
        dept_display = dict(User.UNIT_CHOICES).get(dept, dept)
        writer.writerow([dept, dept_display, 'No Submission'])
    
    writer.writerow([])
    
    # DEPARTMENT RANKING BY BUDGET
    writer.writerow(['=' * 80])
    writer.writerow(['Units Ranking by Budget Allocation'])
    writer.writerow(['=' * 80])
    writer.writerow([])
    
    writer.writerow(['Rank', 'Unit', 'Display Name', 'Total Budget', 'Submissions', 'Average per Submission'])
    for i, dept in enumerate(dept_ranking, 1):
        dept_display = dict(User.UNIT_CHOICES).get(dept['department'], dept['department'])
        dept_requests_count = opb_requests.filter(department=dept['department']).count()
        avg_per_request = dept['total'] / dept_requests_count if dept_requests_count > 0 else 0
        
        writer.writerow([
            i,
            dept['department'],
            dept_display,
            f"₱{dept['total']:,.2f}",
            dept_requests_count,
            f"₱{avg_per_request:,.2f}"
        ])
    
    writer.writerow([])
    
    # DETAILED REQUEST BREAKDOWN
    writer.writerow(['=' * 80])
    writer.writerow(['DETAILED OPB SUBMISSION BREAKDOWN'])
    writer.writerow(['=' * 80])
    writer.writerow([])
    
    writer.writerow([
        'OPB Submission ID', 'Unit Head', 'Unit', 
        'Fiscal Year', 'Total Budget', 'Submissions', 'Status', 
        'Admin Notes', 'Created Date', 'Updated Date'
    ])
    
    for req in opb_requests.order_by('-created_at'):
        writer.writerow([
            str(req.id)[:8] + '...',  # Truncated ID for readability
            f"{req.department_head.first_name} {req.department_head.last_name}",
            req.get_department_display(),
            req.unit or 'N/A',
            req.fiscal_year,
            f"₱{req.total_budget_amount:,.2f}",
            req.item_count,
            req.get_status_display(),
            req.admin_notes or 'N/A',
            req.created_at.strftime('%Y-%m-%d'),
            req.updated_at.strftime('%Y-%m-%d')
        ])
    
    writer.writerow([])
    
    # ITEM-LEVEL DETAILS
    writer.writerow(['=' * 80])
    writer.writerow(['ITEM-LEVEL OPB SUBMISSION DETAILS'])
    writer.writerow(['=' * 80])
    writer.writerow([])
    
    writer.writerow([
        'OPB Submission ID', 'Unit', 'KRA No', 'Objective No', 
        'Indicators', 'Annual Target', 'Activities', 'Timeframe',
        'Budget Amount', 'Source of Fund', 'Responsible Units'
    ])
    
    for req in opb_requests.order_by('-created_at'):
        for item in req.items.all():
            writer.writerow([
                str(req.id)[:8] + '...',
                req.get_department_display(),
                item.kra_no or 'N/A',
                item.objective_no or 'N/A',
                item.indicators or 'N/A',
                item.annual_target or 'N/A',
                item.activities or 'N/A',
                item.timeframe or 'N/A',
                f"₱{item.budget_amount:,.2f}",
                item.source_of_fund or 'N/A',
                item.responsible_units or 'N/A'
            ])
    
    writer.writerow([])
    writer.writerow(['=' * 80])
    writer.writerow(['END OF REPORT'])
    writer.writerow(['=' * 80])
    
    return response


def export_reports_pdf(opb_requests, dept_ranking):
    """Export reports as PDF"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        
        # Create response
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="caobp_report.pdf"'
        
        # Create PDF document
        doc = SimpleDocTemplate(response, pagesize=A4, topMargin=1*inch, bottomMargin=1*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        section_style = ParagraphStyle(
            'SectionTitle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.darkblue,
            borderWidth=1,
            borderColor=colors.darkblue,
            borderPadding=5
        )
        
        # Title and Header
        story.append(Paragraph("CAOBP SYSTEM", title_style))
        story.append(Paragraph("REPORT", title_style))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
        story.append(Paragraph("Generated by: System Administrator", styles['Normal']))
        story.append(Spacer(1, 30))
        
        # Calculate comprehensive statistics
        total_requests = opb_requests.count()
        approved_requests = opb_requests.filter(status='for-approval').count()
        pending_requests = opb_requests.filter(status='pending').count()
        rejected_requests = opb_requests.filter(status='enhancement').count()
        
        total_budget_amount = sum(r.total_budget_amount for r in opb_requests)
        approved_budget_amount = sum(r.total_budget_amount for r in opb_requests.filter(status='for-approval'))
        pending_budget_amount = sum(r.total_budget_amount for r in opb_requests.filter(status='pending'))
        rejected_budget_amount = sum(r.total_budget_amount for r in opb_requests.filter(status='enhancement'))
        
        # Get all departments and calculate submission statistics
        all_departments = User.objects.filter(role='unit_head').values_list('department', flat=True).distinct()
        submitted_departments = set(opb_requests.values_list('department', flat=True))
        not_submitted_departments = set(all_departments) - submitted_departments
        submission_percentage = (len(submitted_departments) / len(all_departments) * 100) if all_departments else 0
        
        # EXECUTIVE SUMMARY
        story.append(Paragraph("OVERALL SUMMARY", section_style))
        
        # Overall Statistics Table
        overall_data = [
            ['Metric', 'Value'],
            ['Total Departments', str(len(all_departments))],
            ['Units with Submissions', str(len(submitted_departments))],
            ['Units without Submissions', str(len(not_submitted_departments))],
            ['Submission Rate', f"{submission_percentage:.1f}%"],
        ]
        
        overall_table = Table(overall_data, colWidths=[2.5*inch, 2*inch])
        overall_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
        ]))
        story.append(overall_table)
        story.append(Spacer(1, 20))
        
        # Request Statistics Table
        request_data = [
            ['Status', 'Count', 'Percentage', 'Budget Amount'],
            ['For Approval', str(approved_requests), f"{(approved_requests/total_requests*100):.1f}%" if total_requests > 0 else "0%", f"₱{approved_budget_amount:,.2f}"],
            ['Pending', str(pending_requests), f"{(pending_requests/total_requests*100):.1f}%" if total_requests > 0 else "0%", f"₱{pending_budget_amount:,.2f}"],
            ['Enhancement', str(rejected_requests), f"{(rejected_requests/total_requests*100):.1f}%" if total_requests > 0 else "0%", f"₱{rejected_budget_amount:,.2f}"],
            ['TOTAL', str(total_requests), '100%', f"₱{total_budget_amount:,.2f}"]
        ]
        
        request_table = Table(request_data, colWidths=[1.5*inch, 1*inch, 1*inch, 1.5*inch])
        request_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, -1), (-1, -1), colors.lightblue),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold')
        ]))
        story.append(request_table)
        story.append(PageBreak())
        
        # DEPARTMENT SUBMISSION ANALYSIS
        story.append(Paragraph("Units SUBMISSION ANALYSIS", section_style))
        
        # Departments with submissions
        story.append(Paragraph("Units with Submissions:", styles['Heading3']))
        dept_submission_data = [['Unit', 'Display Name', 'Submissions', 'Total Budget', 'Status Distribution']]
        
        for dept in submitted_departments:
            dept_requests = opb_requests.filter(department=dept)
            dept_count = dept_requests.count()
            dept_total = sum(r.total_budget_amount for r in dept_requests)
            dept_approved = dept_requests.filter(status='for-approval').count()
            dept_pending = dept_requests.filter(status='pending').count()
            dept_rejected = dept_requests.filter(status='enhancement').count()
            
            dept_display = dict(User.UNIT_CHOICES).get(dept, dept)
            status_dist = f"A:{dept_approved} P:{dept_pending} R:{dept_rejected}"
            
            dept_submission_data.append([
                dept,
                dept_display,
                str(dept_count),
                f"₱{dept_total:,.2f}",
                status_dist
            ])
        
        dept_submission_table = Table(dept_submission_data, colWidths=[1.2*inch, 1.5*inch, 0.8*inch, 1.2*inch, 1.3*inch])
        dept_submission_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
        ]))
        story.append(dept_submission_table)
        story.append(Spacer(1, 20))
        
        # Departments without submissions
        if not_submitted_departments:
            story.append(Paragraph("Units without Submissions:", styles['Heading3']))
            no_submission_data = [['Unit', 'Display Name', 'Status']]
            
            for dept in not_submitted_departments:
                dept_display = dict(User.UNIT_CHOICES).get(dept, dept)
                no_submission_data.append([dept, dept_display, 'No Submission'])
            
            no_submission_table = Table(no_submission_data, colWidths=[2*inch, 2*inch, 2*inch])
            no_submission_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.red),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(no_submission_table)
            story.append(Spacer(1, 20))
        
        story.append(PageBreak())
        
        # DEPARTMENT RANKING
        story.append(Paragraph("Units Ranking by Budget Allocation", section_style))
        
        ranking_data = [['Rank', 'Unit', 'Display Name', 'Total Budget', 'Submissions', 'Average per Submission']]
        
        for i, dept in enumerate(dept_ranking, 1):
            dept_display = dict(User.UNIT_CHOICES).get(dept['department'], dept['department'])
            dept_requests_count = opb_requests.filter(department=dept['department']).count()
            avg_per_request = dept['total'] / dept_requests_count if dept_requests_count > 0 else 0
            
            ranking_data.append([
                str(i),
                dept['department'],
                dept_display,
                f"₱{dept['total']:,.2f}",
                str(dept_requests_count),
                f"₱{avg_per_request:,.2f}"
            ])
        
        ranking_table = Table(ranking_data, colWidths=[0.5*inch, 1.2*inch, 1.5*inch, 1.2*inch, 0.8*inch, 1.3*inch])
        ranking_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
        ]))
        story.append(ranking_table)
        story.append(PageBreak())
        
        # DETAILED REQUEST BREAKDOWN
        story.append(Paragraph("DETAILED OPB SUBMISSION BREAKDOWN", section_style))
        
        # Limit to first 20 requests for PDF readability
        limited_requests = opb_requests.order_by('-created_at')[:20]
        
        detail_data = [['OPB Submission ID', 'Unit Head', 'Unit', 'Budget', 'Status', 'Date']]
        
        for req in limited_requests:
            detail_data.append([
                str(req.id)[:8] + '...',
                f"{req.department_head.first_name} {req.department_head.last_name}",
                req.get_department_display(),
                req.unit or 'N/A',
                f"₱{req.total_budget_amount:,.2f}",
                req.get_status_display(),
                req.created_at.strftime('%Y-%m-%d')
            ])
        
        detail_table = Table(detail_data, colWidths=[1*inch, 1.5*inch, 1*inch, 1*inch, 1*inch, 1*inch, 1*inch])
        detail_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
        ]))
        story.append(detail_table)
        
        if opb_requests.count() > 20:
            story.append(Spacer(1, 10))
            story.append(Paragraph(f"Note: Showing first 20 submissions out of {opb_requests.count()} total submissions.", styles['Normal']))
        
        # Footer
        story.append(Spacer(1, 30))
        story.append(Paragraph("END OF REPORT", ParagraphStyle('Footer', parent=styles['Normal'], alignment=TA_CENTER, fontSize=12, textColor=colors.darkblue)))
        
        # Build PDF
        doc.build(story)
        return response
        
    except ImportError:
        return HttpResponse("PDF export requires ReportLab. Please install it with: pip install reportlab", content_type='text/plain')
    except Exception as e:
        return HttpResponse(f"Error generating PDF: {str(e)}", content_type='text/plain')


def export_reports_docx(opb_requests, dept_ranking):
    """Export reports as DOCX"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn
        
        # Create document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title
        title = doc.add_heading('CAOBP SYSTEM', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('REPORT', 0)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date and metadata
        date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        author_para = doc.add_paragraph("Generated by: System Administrator")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty line
        
        # Calculate comprehensive statistics
        total_requests = opb_requests.count()
        approved_requests = opb_requests.filter(status='for-approval').count()
        pending_requests = opb_requests.filter(status='pending').count()
        rejected_requests = opb_requests.filter(status='enhancement').count()
        
        total_budget_amount = sum(r.total_budget_amount for r in opb_requests)
        approved_budget_amount = sum(r.total_budget_amount for r in opb_requests.filter(status='for-approval'))
        pending_budget_amount = sum(r.total_budget_amount for r in opb_requests.filter(status='pending'))
        rejected_budget_amount = sum(r.total_budget_amount for r in opb_requests.filter(status='enhancement'))
        
        # Get all departments and calculate submission statistics
        all_departments = User.objects.filter(role='unit_head').values_list('department', flat=True).distinct()
        submitted_departments = set(opb_requests.values_list('department', flat=True))
        not_submitted_departments = set(all_departments) - submitted_departments
        submission_percentage = (len(submitted_departments) / len(all_departments) * 100) if all_departments else 0
        
        # EXECUTIVE SUMMARY SECTION
        doc.add_heading('OVERALL SUMMARY', level=1)
        
        # Overall Statistics Table
        doc.add_heading('Overall Statistics', level=2)
        overall_data = [
            ['Metric', 'Value'],
            ['Total Units', str(len(all_departments))],
            ['Units with Submissions', str(len(submitted_departments))],
            ['Units without Submissions', str(len(not_submitted_departments))],
            ['Submission Rate', f"{submission_percentage:.1f}%"],
        ]
        
        overall_table = doc.add_table(rows=len(overall_data), cols=2)
        overall_table.style = 'Table Grid'
        overall_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Style the header row
        header_cells = overall_table.rows[0].cells
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, row_data in enumerate(overall_data):
            row = overall_table.rows[i]
            row.cells[0].text = row_data[0]
            row.cells[1].text = row_data[1]
            if i > 0:  # Not header row
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty line
        
        # Request Statistics Table
        doc.add_heading('OPB SUBMISSION STATISTICS', level=2)
        request_data = [
            ['Status', 'Count', 'Percentage', 'Budget Amount'],
            ['For Approval', str(approved_requests), f"{(approved_requests/total_requests*100):.1f}%" if total_requests > 0 else "0%", f"₱{approved_budget_amount:,.2f}"],
            ['Pending', str(pending_requests), f"{(pending_requests/total_requests*100):.1f}%" if total_requests > 0 else "0%", f"₱{pending_budget_amount:,.2f}"],
            ['Enhancement', str(rejected_requests), f"{(rejected_requests/total_requests*100):.1f}%" if total_requests > 0 else "0%", f"₱{rejected_budget_amount:,.2f}"],
            ['TOTAL', str(total_requests), '100%', f"₱{total_budget_amount:,.2f}"]
        ]
        
        request_table = doc.add_table(rows=len(request_data), cols=4)
        request_table.style = 'Table Grid'
        request_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Style the header and total rows
        for i, row_data in enumerate(request_data):
            row = request_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i == 0 or i == len(request_data) - 1:  # Header or total row
                    row.cells[j].paragraphs[0].runs[0].bold = True
                row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # DEPARTMENT SUBMISSION ANALYSIS
        doc.add_heading('Units SUBMISSION ANALYSIS', level=1)

        # Departments with submissions
        doc.add_heading('Units with Submissions', level=2)
        dept_submission_data = [['Unit', 'Display Name', 'Requests', 'Total Budget', 'Status Distribution']]
        
        for dept in submitted_departments:
            dept_requests = opb_requests.filter(department=dept)
            dept_count = dept_requests.count()
            dept_total = sum(r.total_budget_amount for r in dept_requests)
            dept_approved = dept_requests.filter(status='for-approval').count()
            dept_pending = dept_requests.filter(status='pending').count()
            dept_rejected = dept_requests.filter(status='enhancement').count()
            
            dept_display = dict(User.UNIT_CHOICES).get(dept, dept)
            status_dist = f"A:{dept_approved} P:{dept_pending} R:{dept_rejected}"
            
            dept_submission_data.append([
                dept,
                dept_display,
                str(dept_count),
                f"₱{dept_total:,.2f}",
                status_dist
            ])
        
        dept_submission_table = doc.add_table(rows=len(dept_submission_data), cols=5)
        dept_submission_table.style = 'Table Grid'
        dept_submission_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Style the header row
        header_cells = dept_submission_table.rows[0].cells
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, row_data in enumerate(dept_submission_data):
            row = dept_submission_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i > 0:  # Not header row
                    row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty line
        
        # Departments without submissions
        if not_submitted_departments:
            doc.add_heading('Units without Submissions', level=2)
            no_submission_data = [['Unit', 'Display Name', 'Status']]
            
            for dept in not_submitted_departments:
                dept_display = dict(User.UNIT_CHOICES).get(dept, dept)
                no_submission_data.append([dept, dept_display, 'No Submission'])
            
            no_submission_table = doc.add_table(rows=len(no_submission_data), cols=3)
            no_submission_table.style = 'Table Grid'
            no_submission_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Style the header row
            header_cells = no_submission_table.rows[0].cells
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, row_data in enumerate(no_submission_data):
                row = no_submission_table.rows[i]
                for j, cell_data in enumerate(row_data):
                    row.cells[j].text = cell_data
                    if i > 0:  # Not header row
                        row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # DEPARTMENT RANKING
        doc.add_heading('Units Ranking by Budget Allocation', level=1)
        
        ranking_data = [['Rank', 'Unit', 'Display Name', 'Total Budget', 'Submissions', 'Average per Submission']]
        
        for i, dept in enumerate(dept_ranking, 1):
            dept_display = dict(User.UNIT_CHOICES).get(dept['department'], dept['department'])
            dept_requests_count = opb_requests.filter(department=dept['department']).count()
            avg_per_request = dept['total'] / dept_requests_count if dept_requests_count > 0 else 0
            
            ranking_data.append([
                str(i),
                dept['department'],
                dept_display,
                f"₱{dept['total']:,.2f}",
                str(dept_requests_count),
                f"₱{avg_per_request:,.2f}"
            ])
        
        ranking_table = doc.add_table(rows=len(ranking_data), cols=6)
        ranking_table.style = 'Table Grid'
        ranking_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Style the header row
        header_cells = ranking_table.rows[0].cells
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, row_data in enumerate(ranking_data):
            row = ranking_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i > 0:  # Not header row
                    row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # DETAILED REQUEST BREAKDOWN
        doc.add_heading('DETAILED OPB SUBMISSION BREAKDOWN', level=1)
        
        # Limit to first 30 requests for DOCX readability
        limited_requests = opb_requests.order_by('-created_at')[:30]
        
        detail_data = [['OPB Submission ID', 'Unit Head', 'Unit', 'Budget', 'Status', 'Date']]
        
        for req in limited_requests:
            detail_data.append([
                str(req.id)[:8] + '...',
                f"{req.department_head.first_name} {req.department_head.last_name}",
                req.get_department_display(),
                req.unit or 'N/A',
                f"₱{req.total_budget_amount:,.2f}",
                req.get_status_display(),
                req.created_at.strftime('%Y-%m-%d')
            ])
        
        detail_table = doc.add_table(rows=len(detail_data), cols=7)
        detail_table.style = 'Table Grid'
        detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Style the header row
        header_cells = detail_table.rows[0].cells
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, row_data in enumerate(detail_data):
            row = detail_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i > 0:  # Not header row
                    row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if opb_requests.count() > 30:
            doc.add_paragraph(f"Note: Showing first 30 submissions out of {opb_requests.count()} total submissions.")
        
        # ITEM-LEVEL DETAILS
        doc.add_heading('ITEM-LEVEL OPB SUBMISSION DETAILS', level=1)
        
        # Limit to first 50 items for DOCX readability
        item_count = 0
        item_data = [['OPB Submission ID', 'Unit', 'KRA No', 'Objective No', 'Indicators', 'Budget Amount', 'Source of Fund']]
        
        for req in opb_requests.order_by('-created_at'):
            for item in req.items.all():
                if item_count >= 50:
                    break
                item_data.append([
                    str(req.id)[:8] + '...',
                    req.get_department_display(),
                    item.kra_no or 'N/A',
                    item.objective_no or 'N/A',
                    item.indicators[:50] + '...' if item.indicators and len(item.indicators) > 50 else (item.indicators or 'N/A'),
                    f"₱{item.budget_amount:,.2f}",
                    item.source_of_fund or 'N/A'
                ])
                item_count += 1
            if item_count >= 50:
                break
        
        item_table = doc.add_table(rows=len(item_data), cols=7)
        item_table.style = 'Table Grid'
        item_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Style the header row
        header_cells = item_table.rows[0].cells
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, row_data in enumerate(item_data):
            row = item_table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                if i > 0:  # Not header row
                    row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if item_count >= 50:
            doc.add_paragraph(f"Note: Showing first 50 items out of {sum(req.item_count for req in opb_requests)} total items.")
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph('END OF REPORT')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="caobp_comprehensive_report.docx"'
        
        # Save document to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        response.write(doc_io.getvalue())
        
        return response
        
    except ImportError:
        return HttpResponse("DOCX export requires python-docx. Please install it with: pip install python-docx", content_type='text/plain')
    except Exception as e:
        return HttpResponse(f"Error generating DOCX: {str(e)}", content_type='text/plain')